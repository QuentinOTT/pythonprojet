"""
report.py — Partie 2 : Livre Gutenberg + Image + Rapport Word
"""
import os, re, requests
from collections import Counter
from io import BytesIO
from PIL import Image, ImageOps, ImageDraw
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BOOK_URL   = "https://www.gutenberg.org/cache/epub/1342/pg1342.txt"
IMAGE1_URL = "https://upload.wikimedia.org/wikipedia/commons/thumb/1/15/PrideAndPrejudiceTitlePage.jpg/400px-PrideAndPrejudiceTitlePage.jpg"
REPORTER   = "Quentin Ott"
OUT_DIR    = "rapport_output"
os.makedirs(OUT_DIR, exist_ok=True)


def download_text(url):
    r = requests.get(url, timeout=20)
    r.raise_for_status()
    r.encoding = "utf-8"
    return r.text

def parse_metadata(text):
    title  = re.search(r"^Title:\s*(.+)$",  text, re.MULTILINE)
    author = re.search(r"^Author:\s*(.+)$", text, re.MULTILINE)
    return (
        title.group(1).strip()  if title  else "Pride and Prejudice",
        author.group(1).strip() if author else "Jane Austen"
    )

def extract_first_chapter(text):
    pattern = r"(Chapter I|CHAPTER I|Chapter 1|CHAPTER 1)(?![IV0-9])"
    starts  = [m.start() for m in re.finditer(pattern, text)]
    if not starts:
        raise ValueError("Chapitre 1 introuvable.")
    start       = starts[0]
    end_pattern = r"(Chapter II|CHAPTER II|Chapter 2|CHAPTER 2)(?![IV0-9])"
    ends        = [m.start() for m in re.finditer(end_pattern, text)]
    end         = ends[0] if ends else start + 8000
    return text[start:end].strip()

def split_paragraphs(chapter):
    raw = re.split(r'\n\s*\n', chapter)
    return [p.strip().replace('\n', ' ') for p in raw if len(p.strip()) > 30]

def word_count(text):
    return len(re.findall(r'\b\w+\b', text))

def round_to_tens(n):
    return (n // 10) * 10

def paragraph_stats(paragraphs):
    counts  = [word_count(p) for p in paragraphs]
    rounded = [round_to_tens(c) for c in counts]
    distrib = Counter(sorted(rounded))
    return counts, rounded, distrib

def create_chart(distrib, path):
    xs = sorted(distrib.keys())
    ys = [distrib[x] for x in xs]
    plt.figure(figsize=(11, 6))
    bars = plt.bar(xs, ys, width=8, color="#4472C4", edgecolor="#1F3864")
    for bar, y in zip(bars, ys):
        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                 str(y), ha="center", va="bottom", fontsize=8)
    plt.title("Distribution des longueurs de paragraphes – Chapitre 1",
              fontsize=14, fontweight="bold")
    plt.xlabel("Nombre de mots (arrondi à la dizaine)", fontsize=11)
    plt.ylabel("Nombre de paragraphes", fontsize=11)
    plt.xticks(xs, rotation=45, ha="right")
    plt.grid(axis="y", linestyle="--", alpha=0.4)
    plt.tight_layout()
    plt.savefig(path, dpi=200)
    plt.close()

def download_image(url, path):
    r = requests.get(url, timeout=20)
    r.raise_for_status()
    img = Image.open(BytesIO(r.content)).convert("RGB")
    img.save(path)
    return img

def crop_resize(img, path):
    w, h    = img.size
    cropped = img.crop((int(w * 0.05), int(h * 0.05),
                        int(w * 0.95), int(h * 0.95)))
    resized = cropped.resize((600, 800), Image.LANCZOS)
    resized.save(path)
    return resized

def make_bw_logo(path):
    img  = Image.new("L", (160, 160), 255)
    draw = ImageDraw.Draw(img)
    draw.ellipse([10, 10, 150, 150], outline=0, width=8)
    draw.line([40, 40, 120, 120], fill=0, width=6)
    draw.line([120, 40, 40, 120], fill=0, width=6)
    img.save(path)

def paste_logo(base_path, logo_path, out_path):
    base = Image.open(base_path).convert("RGBA")
    logo = Image.open(logo_path).convert("L")
    logo = ImageOps.autocontrast(logo)
    logo = logo.rotate(30, expand=True, fillcolor=255)
    logo = logo.resize((100, 100), Image.LANCZOS)
    logo_rgba = Image.new("RGBA", logo.size)
    for x in range(logo.width):
        for y in range(logo.height):
            v = logo.getpixel((x, y))
            logo_rgba.putpixel((x, y), (30, 30, 30, max(0, 220 - v)))
    pos = (base.width - logo.width - 15, base.height - logo.height - 15)
    base.alpha_composite(logo_rgba, pos)
    base.convert("RGB").save(out_path)

def build_word_report(title, author, reporter, image_path, chart_path, stats, out_path):
    doc = Document()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(26)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()
    try:
        doc.add_picture(image_path, width=Inches(4.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        doc.add_paragraph("[Image non disponible]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Auteur : {author}")
    r2.italic = True; r2.font.size = Pt(13)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Rapport rédigé par : {reporter}")
    r3.bold = True; r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_page_break()

    # Titre gras italique (style demandé)
    ph = doc.add_paragraph()
    rh = ph.add_run("Distribution des longueurs de paragraphes")
    rh.bold = True; rh.italic = True; rh.font.size = Pt(18)
    rh.font.name = "Times New Roman"
    rh.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    try:
        doc.add_picture(chart_path, width=Inches(6.5))
    except Exception:
        doc.add_paragraph("[Graphique non disponible]")

    doc.add_paragraph()
    ph2 = doc.add_paragraph()
    rh2 = ph2.add_run("Analyse du Chapitre 1")
    rh2.bold = True; rh2.font.size = Pt(15)
    rh2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    desc = (
        f"L'œuvre analysée est « {title} » de {author}, "
        f"téléchargée depuis le Projet Gutenberg.\n\n"
        f"Le premier chapitre contient {stats['num_paragraphs']} paragraphe(s) "
        f"et {stats['total_words']:,} mots au total. "
        f"Le plus court : {stats['min_words']} mots, le plus long : {stats['max_words']} mots. "
        f"Moyenne : {stats['avg_words']:.1f} mots par paragraphe.\n\n"
        f"Le graphique montre la distribution arrondies à la dizaine inférieure "
        f"(ex : 123, 127, 129 → 120)."
    )
    for block in desc.split("\n\n"):
        p = doc.add_paragraph(block.strip())
        p.paragraph_format.space_after = Pt(6)

    ph3 = doc.add_paragraph()
    rh3 = ph3.add_run("Source des données")
    rh3.bold = True; rh3.italic = True; rh3.font.size = Pt(12)
    rh3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph(f"Texte : {BOOK_URL}")
    doc.add_paragraph(f"Image : {IMAGE1_URL}")
    doc.save(out_path)
    print(f"✅ Rapport Word enregistré : {out_path}")


def main():
    print("=== Partie 2 — Génération du rapport ===")
    try:
        print("1. Téléchargement du livre…")
        text = download_text(BOOK_URL)
        print("2. Métadonnées…")
        title, author = parse_metadata(text)
        print(f"   {title} — {author}")
        print("3. Chapitre 1…")
        chap1 = extract_first_chapter(text)
        print("4. Analyse paragraphes…")
        paragraphs = split_paragraphs(chap1)
        counts, rounded, distrib = paragraph_stats(paragraphs)
        stats = {
            "num_paragraphs": len(paragraphs),
            "total_words":    sum(counts),
            "min_words":      min(counts),
            "max_words":      max(counts),
            "avg_words":      sum(counts) / len(counts),
        }
        print(f"   {stats['num_paragraphs']} paragraphes, {stats['total_words']} mots")
        print("5. Graphique…")
        chart_path = os.path.join(OUT_DIR, "chart_distribution.png")
        create_chart(distrib, chart_path)
        print("6. Image #1…")
        raw_path    = os.path.join(OUT_DIR, "image1_raw.jpg")
        edited_path = os.path.join(OUT_DIR, "image1_edited.jpg")
        final_path  = os.path.join(OUT_DIR, "image1_final.jpg")
        img = download_image(IMAGE1_URL, raw_path)
        crop_resize(img, edited_path)
        print("7. Logo BW + assemblage…")
        logo_path = "logo_bw.png"
        if not os.path.exists(logo_path):
            logo_path = os.path.join(OUT_DIR, "logo_bw.png")
            make_bw_logo(logo_path)
        paste_logo(edited_path, logo_path, final_path)
        print("8. Rapport Word…")
        safe     = re.sub(r'[^a-zA-Z0-9_-]', '_', title)
        out_docx = os.path.join(OUT_DIR, f"{safe}_rapport.docx")
        build_word_report(title, author, REPORTER, final_path, chart_path, stats, out_docx)
    except requests.exceptions.ConnectionError:
        print("❌ Erreur réseau.")
    except Exception as exc:
        print(f"❌ Erreur : {exc}")


if __name__ == "__main__":
    main()